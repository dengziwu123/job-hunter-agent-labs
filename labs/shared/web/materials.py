from __future__ import annotations

import io
import json
import re
import shutil
import uuid
from datetime import UTC, datetime
from pathlib import Path
from typing import Any

from labs.shared.config import ROOT_DIR
from labs.shared.web.contracts import JobMaterialPublic, MaterialKind
from labs.shared.web.web_fetch import normalize_web_url


MAX_FILE_BYTES = 5 * 1024 * 1024
MAX_WEB_SOURCES = 10
VALID_ID = re.compile(r"^[A-Za-z0-9_-]{8,100}$")
TEXT_EXTENSIONS = {".txt", ".md"}
SUPPORTED_EXTENSIONS = TEXT_EXTENSIONS | {".json", ".pdf", ".docx"}


class MaterialStore:
    def __init__(self, root: Path | None = None) -> None:
        self.root = root or ROOT_DIR / ".runtime" / "job_materials"

    def create_text(
        self,
        workspace_id: str,
        kind: MaterialKind,
        display_name: str,
        text: str,
        source: str = "paste",
    ) -> JobMaterialPublic:
        workspace = self._workspace(workspace_id)
        normalized = normalize_text(text)
        if not normalized:
            raise ValueError("The material does not contain readable text.")
        if len(normalized.encode("utf-8")) > MAX_FILE_BYTES:
            raise ValueError("The extracted text exceeds the 5 MB limit.")

        existing = self._records(workspace_id)
        if kind == "web_source":
            count = sum(record["kind"] == kind for record in existing)
            if count >= MAX_WEB_SOURCES:
                raise ValueError("A workspace can contain at most 10 Web/Evidence sources.")
        else:
            for record in existing:
                if record["kind"] == kind:
                    self._record_path(workspace, record["material_id"]).unlink(missing_ok=True)

        material_id = f"mat_{uuid.uuid4().hex}"
        record = {
            "material_id": material_id,
            "kind": kind,
            "display_name": safe_display_name(display_name),
            "text": normalized,
            "characters": len(normalized),
            "source": source,
            "status": "ready",
            "source_url": None,
            "created_at": datetime.now(UTC).isoformat(),
        }
        workspace.mkdir(parents=True, exist_ok=True)
        self._record_path(workspace, material_id).write_text(
            json.dumps(record, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        return public_record(record)

    def create_upload(
        self,
        workspace_id: str,
        kind: MaterialKind,
        filename: str,
        content_type: str | None,
        content: bytes,
    ) -> JobMaterialPublic:
        if not content:
            raise ValueError("The uploaded file is empty.")
        if len(content) > MAX_FILE_BYTES:
            raise ValueError("The uploaded file exceeds the 5 MB limit.")
        suffix = Path(filename).suffix.lower()
        if suffix not in SUPPORTED_EXTENSIONS:
            raise ValueError("Supported formats are TXT, Markdown, JSON, PDF, and DOCX.")
        validate_content_type(suffix, content_type)
        text = extract_text(suffix, content)
        return self.create_text(workspace_id, kind, filename, text, source="upload")

    def create_url(self, workspace_id: str, url: str) -> JobMaterialPublic:
        workspace = self._workspace(workspace_id)
        normalized_url = normalize_web_url(url)
        existing = self._records(workspace_id)
        for record in existing:
            if record.get("source_url") == normalized_url:
                return public_record(record)
        count = sum(record["kind"] == "web_source" for record in existing)
        if count >= MAX_WEB_SOURCES:
            raise ValueError("A workspace can contain at most 10 web sources.")

        material_id = f"mat_{uuid.uuid4().hex}"
        host = normalized_url.split("/", 3)[2]
        record = {
            "material_id": material_id,
            "kind": "web_source",
            "display_name": safe_display_name(host),
            "text": "",
            "characters": 0,
            "source": "web",
            "status": "pending",
            "source_url": normalized_url,
            "created_at": datetime.now(UTC).isoformat(),
        }
        workspace.mkdir(parents=True, exist_ok=True)
        self._record_path(workspace, material_id).write_text(
            json.dumps(record, ensure_ascii=False, indent=2) + "\n",
            encoding="utf-8",
        )
        return public_record(record)

    def complete_web_source(
        self,
        workspace_id: str,
        material_id: str,
        *,
        title: str,
        final_url: str,
        text: str,
    ) -> JobMaterialPublic:
        workspace = self._workspace(workspace_id)
        if not VALID_ID.fullmatch(material_id):
            raise ValueError("Invalid material id.")
        path = self._record_path(workspace, material_id)
        if not path.is_file():
            raise FileNotFoundError("Material not found.")
        record = json.loads(path.read_text(encoding="utf-8"))
        if record.get("kind") != "web_source" or record.get("source") != "web":
            raise ValueError("Only pending web sources can be completed by the fetcher.")
        normalized = normalize_text(text)
        if not normalized:
            raise ValueError("The webpage did not contain readable text.")
        record.update(
            {
                "display_name": safe_display_name(title.replace("/", "／").replace("\\", "＼")),
                "text": normalized,
                "characters": len(normalized),
                "status": "ready",
                "source_url": normalize_web_url(final_url),
            }
        )
        path.write_text(json.dumps(record, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        return public_record(record)

    def create_defaults(self, workspace_id: str) -> list[JobMaterialPublic]:
        # The base UI must initialize from the Lab 1 zip by itself. Web
        # research becomes a default only after the Lab 3 package is present.
        existing_kinds = {record["kind"] for record in self._records(workspace_id)}
        fixture_paths = [
            ("candidate_profile", ROOT_DIR / "labs" / "lab_01" / "data" / "profile.json"),
            ("job_description", ROOT_DIR / "labs" / "lab_01" / "data" / "job_description.json"),
        ]
        web_sources = ROOT_DIR / "labs" / "lab_03" / "data" / "sources.json"
        if web_sources.is_file():
            fixture_paths.append(("web_source", web_sources))
        for kind, path in fixture_paths:
            if kind in existing_kinds:
                continue
            self.create_text(
                workspace_id,
                kind,  # type: ignore[arg-type]
                "Synthetic web research sources" if kind == "web_source" else f"Synthetic · {path.name}",
                path.read_text(encoding="utf-8"),
                source="fixture",
            )
            existing_kinds.add(kind)
        return self.list(workspace_id)

    def list(self, workspace_id: str) -> list[JobMaterialPublic]:
        return [public_record(record) for record in self._records(workspace_id)]

    def context(self, workspace_id: str) -> list[dict[str, Any]]:
        return self._records(workspace_id)

    def delete(self, workspace_id: str, material_id: str) -> None:
        workspace = self._workspace(workspace_id)
        if not VALID_ID.fullmatch(material_id):
            raise ValueError("Invalid material id.")
        path = self._record_path(workspace, material_id)
        if not path.is_file():
            raise FileNotFoundError("Material not found.")
        path.unlink()

    def clear(self, workspace_id: str) -> None:
        workspace = self._workspace(workspace_id)
        if not workspace.exists():
            return
        shutil.rmtree(workspace)

    def _records(self, workspace_id: str) -> list[dict[str, Any]]:
        workspace = self._workspace(workspace_id)
        if not workspace.exists():
            return []
        records: list[dict[str, Any]] = []
        for path in workspace.glob("*.json"):
            record = json.loads(path.read_text(encoding="utf-8"))
            # Migrate local workspaces created before the Web source name was
            # made explicit. This keeps hot reload from breaking existing labs.
            if record.get("kind") == "supporting_source":
                record["kind"] = "web_source"
                path.write_text(json.dumps(record, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
            records.append(record)
        return sorted(records, key=lambda item: item["created_at"])

    def _workspace(self, workspace_id: str) -> Path:
        if not VALID_ID.fullmatch(workspace_id):
            raise ValueError("Invalid workspace id.")
        return self.root / workspace_id

    @staticmethod
    def _record_path(workspace: Path, material_id: str) -> Path:
        return workspace / f"{material_id}.json"


def safe_display_name(filename: str) -> str:
    name = Path(filename.replace("\\", "/")).name.strip()
    return (name or "Untitled material")[:120]


def normalize_text(text: str) -> str:
    clean = text.replace("\x00", "").replace("\r\n", "\n").replace("\r", "\n")
    lines = [line.rstrip() for line in clean.splitlines()]
    return "\n".join(lines).strip()


def public_record(record: dict[str, Any]) -> JobMaterialPublic:
    text = record.get("text", "")
    preview = text[:180] + ("…" if len(text) > 180 else "")
    return JobMaterialPublic(
        material_id=record["material_id"],
        kind=record["kind"],
        display_name=record["display_name"],
        characters=record["characters"],
        preview=preview,
        source=record["source"],
        status=record.get("status", "ready"),
        source_url=record.get("source_url"),
    )


def validate_content_type(suffix: str, content_type: str | None) -> None:
    if not content_type or content_type == "application/octet-stream":
        return
    allowed = {
        ".txt": {"text/plain"},
        ".md": {"text/plain", "text/markdown"},
        ".json": {"application/json", "text/json", "text/plain"},
        ".pdf": {"application/pdf"},
        ".docx": {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"},
    }
    if content_type.split(";", 1)[0].lower() not in allowed[suffix]:
        raise ValueError(f"File content type does not match {suffix}.")


def extract_text(suffix: str, content: bytes) -> str:
    if suffix in TEXT_EXTENSIONS:
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError as exc:
            raise ValueError("Text files must use UTF-8 encoding.") from exc
    if suffix == ".json":
        try:
            data = json.loads(content.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError) as exc:
            raise ValueError("The JSON file is not valid UTF-8 JSON.") from exc
        return json.dumps(data, ensure_ascii=False, indent=2)
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader

            reader = PdfReader(io.BytesIO(content))
            return "\n\n".join(page.extract_text() or "" for page in reader.pages)
        except Exception as exc:
            raise ValueError("The PDF could not be parsed.") from exc
    if suffix == ".docx":
        try:
            from docx import Document

            document = Document(io.BytesIO(content))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            table_cells = [cell.text for table in document.tables for row in table.rows for cell in row.cells]
            return "\n".join(paragraphs + table_cells)
        except Exception as exc:
            raise ValueError("The DOCX file could not be parsed.") from exc
    raise ValueError("Unsupported file type.")
